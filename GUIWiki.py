#GUIWiki.py
import wikipedia

#python to docx
from docx import Document
def Wiki(keyword,lang='th'): #function search คำ เเละสร้าง file
    wikipedia.set_lang(lang)

    # summary สำหรับบทความที่สรุปมาเเล้ว
    data = wikipedia.summary(keyword)#.summary เเค่สรุป

    # page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() #สร้างไฟล์ word ใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2) #ดึงข้อมูลคำว่า หมา มาแสดง
    doc.save(keyword + '.docx') #สร้าง file word มาแสดง
    print('file complete')


#เปลี่ยนเป็นภาษไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox #pop up error

GUI = Tk()
GUI.title('program wiki by TNG')
GUI.geometry('400x300')
# config
FONT1 = ('Angsana New',15)

# คำอธิบาย
L = ttk.Label(GUI, text='ค้นหาบทความ',font=FONT1)
L.pack()

# ช่องค้นหาข้อมูล
v_search = StringVar()#กล่องเก็บ Keyword
E1 = ttk.Entry(GUI,textvariable=v_search,font=FONT1,width=40)
E1.pack(pady=10)

# ปุ่ม search

def Search():
    keyword = v_search.get() # .get ดึงข้อมูลเข้ามา (string var)
    try:
        # ลองค้นหาดูว่า result มีไหม
        language = v_radio.get() #th / en / zh
        Wiki(keyword,language)
        messagebox.showinfo('completed','Search completed and save file completed')
    except:
        # หากรันเเล้วมีปัญหา ให้แจ้งเตือน
        messagebox.showwarning('Keyword Error','Please try again')
        
    # print(wikipedia.search(keyword)) # เเสดง คำ search ที่เกี่ยวข้อง
    # result = wikipedia.summary(keyword)
    # print(result)

B1 = ttk.Button(GUI,text='Search',command=Search) #command "ชื่อ function ที่เราจะใช้"
B1.pack(ipadx=20,ipady=10)

#  เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady=20)

v_radio = StringVar() # ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='อังกฤษ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='จีน',variable=v_radio,value='zh')
RB1.invoke() #set ค่าเริ่มต้นเป็นภาษาไทย

# เรียงใหม่ให้เป็น แนวนอน
RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)

GUI.mainloop()
